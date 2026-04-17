"""
报告生成工具
生成Word格式的分析报告，包含案件概况、可疑线索、证据引用、金额统计、关系图谱
"""
from datetime import datetime
from typing import Optional, Dict, Any, List
import os

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class ReportGenerator:
    """Word报告生成器"""

    # 报告模板目录
    TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), 'templates')

    @classmethod
    def generate_case_report(
        cls,
        case_data: Dict[str, Any],
        chain_analysis: Dict[str, Any],
        evidence_inventory: Dict[str, Any],
        output_path: Optional[str] = None
    ) -> str:
        """
        生成案件分析Word报告

        Args:
            case_data: 案件数据（包含case_detail信息）
            chain_analysis: 上下游关系分析数据
            evidence_inventory: 证据清单数据
            output_path: 输出文件路径，默认为None则自动生成

        Returns:
            报告文件路径
        """
        if not HAS_DOCX:
            # 降级为文本报告
            return cls._generate_text_report(case_data, chain_analysis, evidence_inventory, output_path)

        case = case_data.get("case", {})
        stats = case_data.get("statistics", {})

        doc = Document()

        # 设置文档标题
        title = doc.add_heading('火眼智擎—汽配领域知产保护分析报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 案件基本信息
        doc.add_heading('一、案件概况', 1)
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        table_data = [
            ("案件编号", case.get('case_no', 'N/A')),
            ("嫌疑人姓名", case.get('suspect_name', 'N/A')),
            ("涉案品牌", case.get('brand', 'N/A')),
            ("涉案金额", f"{case.get('amount', 0)}元"),
            ("生成时间", datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        ]
        for i, (key, value) in enumerate(table_data):
            table.rows[i].cells[0].text = key
            table.rows[i].cells[1].text = str(value)

        # 统计信息
        doc.add_heading('二、数据统计', 1)
        stats_table = doc.add_table(rows=5, cols=2)
        stats_table.style = 'Table Grid'
        stats_data = [
            ("交易记录", f"{stats.get('transaction_count', 0)}条"),
            ("通讯记录", f"{stats.get('communication_count', 0)}条"),
            ("物流记录", f"{stats.get('logistics_count', 0)}条"),
            ("涉案人员", f"{stats.get('person_count', 0)}人"),
            ("可疑线索", f"{stats.get('suspicious_clue_count', 0)}条")
        ]
        for i, (key, value) in enumerate(stats_data):
            stats_table.rows[i].cells[0].text = key
            stats_table.rows[i].cells[1].text = value

        # 可疑线索
        doc.add_heading('三、可疑线索', 1)
        suspicious_clues = case_data.get("suspicious_clues", [])
        if suspicious_clues:
            for i, clue in enumerate(suspicious_clues[:10], 1):
                p = doc.add_paragraph()
                p.add_run(f"{i}. ").bold = True
                p.add_run(f"[{clue.get('clue_type', '未知')}] ")
                p.add_run(clue.get('evidence_text', '')[:100])
                p = doc.add_paragraph()
                p.add_run(f"   评分: {clue.get('score', 0)}分 | 涉嫌罪名: {clue.get('crime_type', '待定')}")
        else:
            doc.add_paragraph("暂无可疑线索")

        # 上下游关系
        doc.add_heading('四、上下游关系分析', 1)
        p = doc.add_paragraph()
        p.add_run(f"上游供货商: {len(chain_analysis.get('upstream', []))}个")
        p = doc.add_paragraph()
        p.add_run(f"下游买家: {len(chain_analysis.get('downstream', []))}个")
        p = doc.add_paragraph()
        p.add_run(f"核心嫌疑人: {len(chain_analysis.get('core_suspects', []))}个")

        # 角色分析
        role_analysis = chain_analysis.get("role_analysis", {})
        if role_analysis.get("producers"):
            p = doc.add_paragraph()
            p.add_run("生产者: ").bold = True
            p.add_run(", ".join(role_analysis['producers']))
        if role_analysis.get("sellers"):
            p = doc.add_paragraph()
            p.add_run("销售者: ").bold = True
            p.add_run(", ".join(role_analysis['sellers']))

        # 证据清单
        doc.add_heading('五、证据清单', 1)
        p = doc.add_paragraph()
        p.add_run(f"通讯证据: {evidence_inventory.get('communication_evidence_count', 0)}条")
        p = doc.add_paragraph()
        p.add_run(f"价格异常证据: {evidence_inventory.get('price_anomaly_evidence_count', 0)}条")
        p = doc.add_paragraph()
        p.add_run(f"物流异常证据: {evidence_inventory.get('logistics_evidence_count', 0)}条")

        # 金额统计
        doc.add_heading('六、金额统计', 1)
        amount_summary = case_data.get("amount_summary", {})
        if amount_summary:
            amount_table = doc.add_table(rows=4, cols=2)
            amount_table.style = 'Table Grid'
            amount_data = [
                ("非法经营数额", f"{amount_summary.get('illegal_business_amount', 0)}元"),
                ("违法所得数额", f"{amount_summary.get('illegal_gain_amount', 0)}元"),
                ("刑事门槛", f"{amount_summary.get('criminal_threshold', 'N/A')}"),
                ("是否符合立案标准", "是" if amount_summary.get('meets_criminal_threshold', False) else "否")
            ]
            for i, (key, value) in enumerate(amount_data):
                amount_table.rows[i].cells[0].text = key
                amount_table.rows[i].cells[1].text = value
        else:
            doc.add_paragraph("暂无金额统计数据")

        # 保存文档
        if output_path is None:
            output_dir = os.path.join(os.path.dirname(__file__), '..', 'data', 'reports')
            os.makedirs(output_dir, exist_ok=True)
            filename = f"report_{case.get('case_no', 'unknown')}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
            output_path = os.path.join(output_dir, filename)

        doc.save(output_path)
        return output_path

    @classmethod
    def _generate_text_report(
        cls,
        case_data: Dict[str, Any],
        chain_analysis: Dict[str, Any],
        evidence_inventory: Dict[str, Any],
        output_path: Optional[str] = None
    ) -> str:
        """
        生成文本格式报告（当python-docx未安装时的降级方案）
        """
        case = case_data.get("case", {})
        stats = case_data.get("statistics", {})

        lines = [
            "=" * 60,
            "火眼智擎—汽配领域知产保护分析报告",
            "=" * 60,
            "",
            f"案件编号: {case.get('case_no', 'N/A')}",
            f"嫌疑人姓名: {case.get('suspect_name', 'N/A')}",
            f"涉案品牌: {case.get('brand', 'N/A')}",
            f"涉案金额: {case.get('amount', 0)}元",
            f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            "",
            "-" * 60,
            "一、数据统计",
            "-" * 60,
            f"  交易记录: {stats.get('transaction_count', 0)}条",
            f"  通讯记录: {stats.get('communication_count', 0)}条",
            f"  物流记录: {stats.get('logistics_count', 0)}条",
            f"  涉案人员: {stats.get('person_count', 0)}人",
            f"  可疑线索: {stats.get('suspicious_clue_count', 0)}条",
            "",
            "-" * 60,
            "二、可疑线索",
            "-" * 60,
        ]

        suspicious_clues = case_data.get("suspicious_clues", [])
        if suspicious_clues:
            for i, clue in enumerate(suspicious_clues[:10], 1):
                lines.append(f"  {i}. [{clue.get('clue_type', '未知')}] {clue.get('evidence_text', '')[:50]}...")
                lines.append(f"     评分: {clue.get('score', 0)}分 | 涉嫌罪名: {clue.get('crime_type', '待定')}")
        else:
            lines.append("  暂无可疑线索")

        lines.extend([
            "",
            "-" * 60,
            "三、上下游关系",
            "-" * 60,
            f"  上游供货商: {len(chain_analysis.get('upstream', []))}个",
            f"  下游买家: {len(chain_analysis.get('downstream', []))}个",
            f"  核心嫌疑人: {len(chain_analysis.get('core_suspects', []))}个",
            "",
            "-" * 60,
            "四、证据清单",
            "-" * 60,
            f"  通讯证据: {evidence_inventory.get('communication_evidence_count', 0)}条",
            f"  价格异常证据: {evidence_inventory.get('price_anomaly_evidence_count', 0)}条",
            f"  物流异常证据: {evidence_inventory.get('logistics_evidence_count', 0)}条",
            "",
            "=" * 60,
            "报告生成完毕",
            "=" * 60,
        ])

        content = "\n".join(lines)

        if output_path is None:
            output_dir = os.path.join(os.path.dirname(__file__), '..', 'data', 'reports')
            os.makedirs(output_dir, exist_ok=True)
            filename = f"report_{case.get('case_no', 'unknown')}_{datetime.now().strftime('%Y%m%d%H%M%S')}.txt"
            output_path = os.path.join(output_dir, filename)

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)

        return output_path

    @classmethod
    def generate_simple_report(
        cls,
        case_no: str,
        suspect_name: str,
        brand: str,
        amount: float,
        clue_count: int,
        evidence_count: int,
        output_path: Optional[str] = None
    ) -> str:
        """
        生成简化报告（用于快速生成）

        Args:
            case_no: 案件编号
            suspect_name: 嫌疑人姓名
            brand: 涉案品牌
            amount: 涉案金额
            clue_count: 可疑线索数
            evidence_count: 证据数量
            output_path: 输出路径

        Returns:
            报告文件路径
        """
        case_data = {
            "case": {
                "case_no": case_no,
                "suspect_name": suspect_name,
                "brand": brand,
                "amount": amount
            },
            "statistics": {
                "suspicious_clue_count": clue_count
            },
            "suspicious_clues": []
        }
        chain_analysis = {"upstream": [], "downstream": [], "core_suspects": [], "role_analysis": {}}
        evidence_inventory = {"communication_evidence_count": evidence_count, "price_anomaly_evidence_count": 0, "logistics_evidence_count": 0}

        return cls.generate_case_report(case_data, chain_analysis, evidence_inventory, output_path)
